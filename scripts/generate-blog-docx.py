"""Generate a formatted Word document from the blog markdown."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BLOG_PATH = Path(__file__).parent.parent / "docs" / "blog-agentdedup-architecture.md"
OUTPUT_PATH = Path(__file__).parent.parent / "docs" / "AgentDedup-Blog-Architecture.docx"
HEADER_IMG = Path(__file__).parent.parent / "docs" / "blog-header-agentdedup.png"


def set_style(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x23, 0x2F, 0x3E)
    return h


def add_table_from_lines(doc, header_line, rows):
    """Add a formatted table to the document."""
    headers = [c.strip() for c in header_line.split("|") if c.strip()]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_line in rows:
        cells = [c.strip() for c in row_line.split("|") if c.strip()]
        row = table.add_row()
        for i, val in enumerate(cells):
            if i < len(headers):
                row.cells[i].text = val
                for p in row.cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()


def main():
    doc = Document()
    set_style(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(BLOG_PATH, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_header = ""
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip image references (we'll add the header image manually)
        if line.startswith("!["):
            # Add header image if it's the first one
            if HEADER_IMG.exists() and "header" in line.lower():
                doc.add_picture(str(HEADER_IMG), width=Inches(6))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "Figure" in lines[i + 1] if i + 1 < len(lines) else "":
                # Skip diagram references (draw.io files can't be embedded)
                pass
            i += 1
            continue

        # Figure captions
        if line.startswith("*Figure"):
            p = doc.add_paragraph(line.strip("*"))
            p.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            add_heading(doc, line[2:], 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], 2)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
            i += 1
            continue

        # Table detection
        if "|" in line and "---" not in line:
            if not in_table:
                # Check if next line is separator
                if i + 1 < len(lines) and "---" in lines[i + 1]:
                    in_table = True
                    table_header = line
                    table_rows = []
                    i += 2  # Skip header and separator
                    continue
                else:
                    # Regular line with pipe
                    pass
            else:
                # Collecting table rows
                if line.strip() and "|" in line:
                    table_rows.append(line)
                    i += 1
                    continue
                else:
                    # End of table
                    add_table_from_lines(doc, table_header, table_rows)
                    in_table = False
                    table_header = ""
                    table_rows = []
                    # Don't increment i, process current line
                    continue

        # If we were in a table and hit a non-table line
        if in_table:
            add_table_from_lines(doc, table_header, table_rows)
            in_table = False
            table_header = ""
            table_rows = []

        # Numbered lists
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            text = text.replace("**", "")
            p = doc.add_paragraph(text, style="List Number")
            i += 1
            continue

        # Bullet points
        if line.startswith("- **"):
            text = line[2:].replace("**", "")
            p = doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Bold paragraphs (like talking points)
        if line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            # Handle inline bold
            text = line
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Empty line
        i += 1

    # Flush any remaining table
    if in_table:
        add_table_from_lines(doc, table_header, table_rows)

    # Save
    doc.save(str(OUTPUT_PATH))
    print(f"Word document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
